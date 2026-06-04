from __future__ import annotations

import json
import re
import textwrap
import xml.etree.ElementTree as ET
from pathlib import Path
from html import unescape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ARTIFACT_ROOT = ROOT / "docs" / "grant-review" / "evidence-artifacts"
API_DIR = ARTIFACT_ROOT / "api"
SCREENSHOT_DIR = ARTIFACT_ROOT / "screenshots"
SVG_DIR = ARTIFACT_ROOT / "generated-images"
PNG_CARD_DIR = ARTIFACT_ROOT / "generated-images-png"
OUTPUT = ROOT / "docs" / "grant-review" / "forg3t-avalanche-phase2-evidence-artifacts.docx"


LIVE_LINKS = [
    ("Public verify page", "https://buildgames.forg3t.io/verify/212a028eabcac5e2c3458cc9f219259d"),
    (
        "Public verify API",
        "https://xewxfsdrtqpthkpbhbzp.supabase.co/functions/v1/verify-evidence?token=212a028eabcac5e2c3458cc9f219259d",
    ),
    (
        "Snowtrace transaction",
        "https://snowtrace.io/tx/0x7af8b0376079571f2a4ff46ff76e6cdfb27f710ea4b2434c41bfdf25a167e7be",
    ),
    ("Snowtrace contract", "https://snowtrace.io/address/0x20E772a60CEE7D8E6706E698B129FD917c3936bf"),
]


CANONICAL_IDS = [
    ("Project ID", "0c7643e1-471f-4b04-848c-329c39f77143"),
    ("Job ID", "9a2a77cf-4b09-4f59-be04-c18b08b137bd"),
    ("Evidence ID", "608a427f-25c1-4f43-b3e7-d9a86ff33801"),
    ("Public verify token", "212a028eabcac5e2c3458cc9f219259d"),
    ("Avalanche network", "mainnet"),
    ("Chain ID", "43114"),
    ("Block number", "86276105"),
    ("Contract address", "0x20E772a60CEE7D8E6706E698B129FD917c3936bf"),
]


ARTIFACT_INDEX = [
    ("Live Avalanche transaction link", "Done", "generated-images/01-live-avalanche-transaction.svg", "screenshots/snowtrace-transaction.png, api/anchor-status.json"),
    ("Explorer screenshot", "Done", "screenshots/snowtrace-transaction.png", "screenshots/snowtrace-contract.png"),
    ("Job detail evidence", "Done", "generated-images/02-job-detail-proof.svg", "api/job-detail.json"),
    ("Evidence detail evidence", "Done", "generated-images/03-evidence-detail-proof.svg", "api/authenticated-verify-evidence.json"),
    ("Public verify page screenshot", "Done", "screenshots/public-verify-page.png", "generated-images/04-public-verify-proof.svg"),
    ("Drag and drop verification base UI", "Done", "screenshots/public-verify-page.png", "generated-images/05-drag-drop-json-valid.svg"),
    ("JSON valid verification", "Done", "generated-images/05-drag-drop-json-valid.svg", "exports/forg3t-evidence-bundle.valid.json, api/verify-upload-json-valid-response.json"),
    ("JSON mismatch verification", "Done", "generated-images/06-drag-drop-json-mismatch.svg", "exports/forg3t-evidence-bundle.mismatch.json, api/verify-upload-json-mismatch-response.json"),
    ("PDF verification", "Done", "generated-images/07-drag-drop-pdf-valid.svg", "exports/forg3t-evidence-export.pdf, api/verify-upload-pdf-valid-response.json"),
    ("Unsupported file state", "Done", "generated-images/08-drag-drop-unsupported-file.svg", "api/verify-upload-unsupported-file-response.json"),
    ("JSON export", "Done", "exports/forg3t-evidence-export.json", "generated-images/09-json-export.svg, api/report-export-json-response.json"),
    ("CSV export", "Done", "exports/forg3t-evidence-export.csv", "generated-images/10-csv-export.svg, api/report-export-csv-response.json"),
    ("PDF export", "Done", "exports/forg3t-evidence-export.pdf", "generated-images/11-pdf-export.svg, api/report-pdf-hash-commit.json"),
    ("Pipeline run screenshot/proof", "Done", "generated-images/12-pipeline-run.svg", "api/pipeline-run.json, api/pipeline-detail-with-runs.json"),
    ("Admin role proof", "Done", "generated-images/13-admin-role-proof.svg", "api/project-access-memberships.json"),
    ("SDK/API smoke output", "Done", "generated-images/14-sdk-smoke-output.svg", "logs/phase2-smoke.log"),
    ("Build, test, audit logs", "Done", "generated-images/15-build-test-audit-logs.svg", "logs/npm-test.log, logs/npm-build.log, logs/npm-audit.log, logs/contracts-compile.log, logs/contracts-audit.log"),
    ("Generic HTTP integration smoke", "Done", "generated-images/16-generic-http-integration-smoke.svg", "api/integration-create-generic-http.json, api/integration-test-generic-http.json"),
    ("Production dashboard bundle check", "Resolved", "generated-images/17-production-login-env-risk.svg", "api/production-bundle-env-check.json"),
    ("Pilot evidence memo", "Founder Evidence Required", "generated-images/18-founder-pilot-evidence-required.svg", "Founder must provide external memo if claiming pilot progress."),
    ("Customer or partner confirmation", "Founder Evidence Required", "generated-images/19-customer-confirmation-required.svg", "Founder must provide external confirmation if available."),
    ("Recorded demo video", "Founder Evidence Required", "generated-images/20-demo-video-required.svg", "Founder must record or provide demo link separately."),
]


def read_json(name: str) -> dict:
    path = API_DIR / name
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf") if bold else Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def render_svg_cards() -> list[Path]:
    PNG_CARD_DIR.mkdir(parents=True, exist_ok=True)
    rendered: list[Path] = []
    for svg_path in sorted(SVG_DIR.glob("*.svg")):
        svg = svg_path.read_text(encoding="utf-8", errors="replace")
        svg = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", svg).replace("\ufffd", "")
        width_match = re.search(r'width="([0-9.]+)"', svg)
        height_match = re.search(r'height="([0-9.]+)"', svg)
        width = int(float(width_match.group(1))) if width_match else 1400
        height = int(float(height_match.group(1))) if height_match else 520
        image = Image.new("RGB", (width, height), "#f8fafc")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((28, 28, width - 28, height - 28), radius=18, fill="#ffffff", outline="#d7dee8", width=2)

        text_matches = re.findall(r'<text\s+([^>]+)>(.*?)</text>', svg, flags=re.DOTALL)
        for attr_text, raw_text in text_matches:
            text = unescape(re.sub(r"<[^>]+>", "", raw_text)).strip()
            if not text:
                continue
            x_match = re.search(r'x="([0-9.]+)"', attr_text)
            y_match = re.search(r'y="([0-9.]+)"', attr_text)
            class_match = re.search(r'class="([^"]*)"', attr_text)
            x = int(float(x_match.group(1))) if x_match else 46
            y = int(float(y_match.group(1))) if y_match else 46
            css_class = class_match.group(1) if class_match else ""
            if css_class == "title":
                draw.text((x, y - 31), text, font=font(32, True), fill="#101828")
            elif css_class == "eyebrow":
                draw.text((x, y - 16), text, font=font(17, True), fill="#2f80ed")
            elif css_class == "status":
                # Draw badge background behind status.
                fill = "#fef2f2" if "Risk" in text else "#fff7ed" if "Founder" in text else "#ecfdf5"
                outline = "#fca5a5" if "Risk" in text else "#fdba74" if "Founder" in text else "#86efac"
                draw.rounded_rectangle((1088, 58, 1330, 102), radius=10, fill=fill, outline=outline, width=2)
                draw.text((x, y - 18), text, font=font(18, True), fill="#101828")
            elif css_class == "label":
                draw.text((x, y - 18), text, font=font(18, True), fill="#475467")
            elif css_class == "footer":
                draw.text((x, y - 15), text, font=font(16), fill="#667085")
            else:
                draw.text((x, y - 19), text, font=font(20), fill="#111827")

        output = PNG_CARD_DIR / f"{svg_path.stem}.png"
        image.save(output, "PNG")
        rendered.append(output)
    return rendered


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "DADCE0") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F80ED")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5 if level < 3 else 0x78)
    return paragraph


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (1.85, 4.65)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Field", bold=True, size=9)
    set_cell_text(hdr[1], "Value", bold=True, size=9)
    set_cell_shading(hdr[0], "F2F4F7")
    set_cell_shading(hdr[1], "F2F4F7")
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=8)
        set_cell_text(cells[1], value, size=8)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        row.cells[0].width = Inches(widths[0])
        row.cells[1].width = Inches(widths[1])
    doc.add_paragraph()
    return table


def add_matrix_table(doc: Document):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    headers = ["Checklist Item", "Status", "Primary Artifact", "Supporting Evidence"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=8)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    set_repeat_table_header(table.rows[0])
    for item, status, primary, support in ARTIFACT_INDEX:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, size=7)
        color = "9B1C1C" if status == "Risk noted" else "7A5A00" if status == "Founder Evidence Required" else "166534"
        set_cell_text(cells[1], status, bold=True, color=color, size=7)
        set_cell_text(cells[2], primary, size=7)
        set_cell_text(cells[3], support, size=7)
    widths = [1.62, 0.9, 1.95, 2.03]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def add_note(doc: Document, title: str, body: str, fill: str = "FFF7ED"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, "FDBA74" if fill == "FFF7ED" else "BFDBFE")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.35):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Forg3t Protocol Avalanche Phase 2 Evidence Artifacts")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x70, 0x85)


def main():
    rendered_cards = render_svg_cards()

    doc = Document()
    style_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Forg3t Protocol Avalanche Phase 2 Evidence Artifacts")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph("Reviewer-ready DOCX package generated from repository evidence, live Edge Function responses, exports, and screenshots. Generated 2026-05-25.")
    subtitle.paragraph_format.space_after = Pt(12)

    add_note(
        doc,
        "Boundary",
        "This document does not claim enterprise pilot completion, customer attestation, legal approval, production deployment beyond verified public routes, or recorded demo completion. Founder-only evidence requirements are marked explicitly.",
        "EFF6FF",
    )

    add_heading(doc, "Key Live Links", 1)
    for label, url in LIVE_LINKS:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        add_hyperlink(p, url, url)

    add_heading(doc, "Canonical Evidence IDs", 1)
    add_key_value_table(doc, CANONICAL_IDS)

    add_heading(doc, "Checklist Matrix", 1)
    doc.add_paragraph("Each checklist item below maps to a generated proof card, screenshot, API response, export file, log file, or Founder Evidence Required boundary.")
    add_matrix_table(doc)

    add_heading(doc, "Production Bundle Check", 1)
    prod = read_json("production-bundle-env-check.json")
    risk_text = "The public verify route and Edge Function proof work. The production bundle check detected a service-role-looking JWT string in the currently fetched production JavaScript bundle. Redeploy the dashboard with only the anon key before asking a reviewer to log in through the production dashboard."
    fill = "FFF7ED"
    if not prod.get("serviceRoleJwtStringPresentInBundle"):
        risk_text = "No service-role-looking JWT string was detected in the checked production bundle artifact."
        fill = "ECFDF5"
    add_note(doc, "Production Bundle Check", risk_text, fill)

    add_heading(doc, "Primary Screenshots", 1)
    screenshots = [
        ("public-verify-page.png", "Public verify page showing confirmed anchor state."),
        ("snowtrace-transaction.png", "Snowtrace transaction page for the Avalanche mainnet anchor."),
        ("snowtrace-contract.png", "Snowtrace contract page for the evidence anchor contract."),
        ("local-dashboard-overview.png", "Local dashboard overview using the reviewer account and live Supabase backend."),
    ]
    for filename, caption in screenshots:
        add_image(doc, SCREENSHOT_DIR / filename, caption)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Generated Proof Cards", 1)
    doc.add_paragraph("The following proof cards summarize each checklist item. Source SVG files remain in generated-images; DOCX-embedded PNG renders are in generated-images-png.")
    for card in rendered_cards:
        add_image(doc, card, card.stem.replace("-", " ").title(), width=6.2)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Export And Log Artifacts", 1)
    add_key_value_table(
        doc,
        [
            ("JSON export", "docs/grant-review/evidence-artifacts/exports/forg3t-evidence-export.json"),
            ("CSV export", "docs/grant-review/evidence-artifacts/exports/forg3t-evidence-export.csv"),
            ("PDF export", "docs/grant-review/evidence-artifacts/exports/forg3t-evidence-export.pdf"),
            ("Phase 2 smoke log", "docs/grant-review/evidence-artifacts/logs/phase2-smoke.log"),
            ("npm test log", "docs/grant-review/evidence-artifacts/logs/npm-test.log"),
            ("npm build log", "docs/grant-review/evidence-artifacts/logs/npm-build.log"),
            ("npm audit log", "docs/grant-review/evidence-artifacts/logs/npm-audit.log"),
            ("Contract compile log", "docs/grant-review/evidence-artifacts/logs/contracts-compile.log"),
            ("Contract audit log", "docs/grant-review/evidence-artifacts/logs/contracts-audit.log"),
        ],
    )

    add_heading(doc, "Commands Used", 1)
    for command in [
        "npm test",
        "npm run lint",
        "npm run build",
        "npm audit",
        "cd contracts && npm run compile",
        "cd contracts && npm audit",
        "npm run smoke:phase2",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    add_heading(doc, "Founder Evidence Required", 1)
    doc.add_paragraph("The repository cannot independently provide external human evidence. Before submitting a final grant review package, the founder should provide these items if they are being claimed:")
    for item in [
        "Pilot evidence memo",
        "Customer or partner confirmation, if available",
        "Recorded demo or walkthrough video, if requested by the grant manager",
        "Any enterprise approval, customer usage proof, or sales/procurement material that is not present in the repository",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
