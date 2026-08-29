from pathlib import Path
from zipfile import ZipFile
import json

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "grant-review"
ASSET_DIR = OUT_DIR / "submission-assets"
SS_DIR = ASSET_DIR / "screenshots"
VIDEO_DIR = ASSET_DIR / "video"
API_DIR = ASSET_DIR / "api"
RENDER_DIR = ASSET_DIR / "renders"

DOCX_PATH = OUT_DIR / "forg3t-avalanche-phase2-submission-packet.docx"
GIF_PATH = VIDEO_DIR / "forg3t-avalanche-phase2-proof-walkthrough.gif"
THUMB_PATH = VIDEO_DIR / "forg3t-avalanche-phase2-proof-walkthrough-cover.png"

INK = "111111"
MUTED = "4B5563"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BORDER = "DADCE0"
LIGHT_FILL = "F2F4F7"
GREEN = "0F766E"

REPO_URL = "https://github.com/Alvinagile/Forg3t-Avax"
LIVE_APP = "https://buildgames.forg3t.io"
MEDIUM_ARTICLE = "https://medium.com/@aiunlearning/forg3t-protocol-and-avalanche-evidence-anchoring-7310bf22226e"
PILOT_CONTRACT_LINK = "https://drive.google.com/file/d/1XHKNlS2EuiAAWFcff0P-Wgo-uuABiNUg/view?usp=sharing"
LATEST_VERIFY = "https://buildgames.forg3t.io/verify/636919b6e97441df953d98278bbc0efe"
LATEST_TX = "https://snowtrace.io/tx/0x4f9d253d097808406777ba1b9d67b0ac4baac44d34d1d101cea4a3721559c69b"
CONTRACT = "0x20E772a60CEE7D8E6706E698B129FD917c3936bf"
LATEST_TX_HASH = "0x4f9d253d097808406777ba1b9d67b0ac4baac44d34d1d101cea4a3721559c69b"
LATEST_BLOCK = "87016942"
PROJECT_ID = "0c7643e1-471f-4b04-848c-329c39f77143"

PREVIOUS_MAINNET_ANCHORS = [
    {
        "date": "2026-06-02",
        "job": "dcef5dc1-3b85-4e36-8ab8-55f0cd8ea803",
        "evidence": "308463f6-cad9-4490-ac9b-af4ac3cfd8b1",
        "tx": LATEST_TX_HASH,
        "block": LATEST_BLOCK,
    },
    {
        "date": "2026-05-31",
        "job": "c4fc546a-b544-4d27-a8d3-abba1118f1d9",
        "evidence": "180594c5-17e4-48e3-98c2-56965b8964b9",
        "tx": "0x1d73a8ef8028c81292caf5b9081a9ff4d4d47519f24008abc5fa6f7d6b56453b",
        "block": "mainnet",
    },
    {
        "date": "2026-05-29",
        "job": "c54e7f03-16dc-4a3d-a902-2d4c9ed91005",
        "evidence": "fc8a2a5e-c14c-48ac-b49d-abce4e57adc8",
        "tx": "0x43614297bb117278dbd481eb615fff8649cff7bbc3c87c63c6c9691855017356",
        "block": "mainnet",
    },
    {
        "date": "2026-05-28",
        "job": "84c6009a-a9bd-4a21-82d5-a92920fb06fe",
        "evidence": "3561358e-8944-4bab-9f1d-5a32007f6ab6",
        "tx": "0xf5edaff5777f349ac3d1cb00bfc5953040927b74450667194a00cd7d844678dc",
        "block": "mainnet",
    },
    {
        "date": "2026-05-27",
        "job": "daily-reviewer-anchor-run",
        "evidence": "reviewer automation evidence",
        "tx": "0x751fec0d93290cc0bcc25e42735069e88c42dd42c664e6ed064b025c22e1ec8d",
        "block": "mainnet",
    },
]

MILESTONES = [
    {
        "number": 1,
        "name": "Launch production grade evidence anchoring on Avalanche",
        "proof": "Avalanche C-Chain mainnet transaction, contract bytecode, receipt success, block number, and Snowtrace proof.",
        "screenshot": "milestone-01-avalanche-mainnet-anchor.png",
        "links": [("Snowtrace transaction", LATEST_TX)],
        "files": [
            "project/supabase/functions/anchors/index.ts",
            "project/supabase/functions/verify-evidence/index.ts",
            "project/docs/grant-review/submission-assets/api/avalanche-mainnet-verification.json",
        ],
    },
    {
        "number": 2,
        "name": "Ship drag and drop verification flow for auditors and third parties",
        "proof": "Public verification route supports JSON/PDF upload, anchor-confirmed state, local hash comparison, and auditor-facing result states.",
        "screenshot": "milestone-02-drag-drop-verification.png",
        "links": [("Public verify route", LATEST_VERIFY)],
        "files": [
            "project/src/pages/Verify.tsx",
            "project/supabase/functions/verify-evidence/index.ts",
            "project/docs/grant-review/evidence-artifacts/api/verify-upload-json-valid-response.json",
            "project/docs/grant-review/evidence-artifacts/api/verify-upload-pdf-valid-response.json",
        ],
    },
    {
        "number": 3,
        "name": "Complete job history, transaction visibility, and evidence verification UX",
        "proof": "Populated reviewer workspace shows completed jobs, evidence readiness, confirmed anchor state, and transaction hash visibility.",
        "screenshot": "milestone-03-job-history-populated.png",
        "links": [("Live app", LIVE_APP)],
        "files": [
            "project/src/pages/Jobs.tsx",
            "project/src/pages/JobDetail.tsx",
            "project/src/pages/EvidenceDetail.tsx",
            "project/docs/grant-review/submission-assets/screenshots/milestone-03-job-history-populated.png",
        ],
    },
    {
        "number": 4,
        "name": "Publish technical architecture and compliance workflow documentation",
        "proof": "Public Medium article and repository docs explain Forg3t, Avalanche anchoring, compliance workflow, API lifecycle, and reviewer runbook.",
        "screenshot": "milestone-04-medium-architecture-article.png",
        "links": [("Medium article", MEDIUM_ARTICLE)],
        "files": [
            "project/docs/phase2-readiness.md",
            "project/docs/avalanche/technicalArchitecture.md",
            "project/docs/avalanche/reviewerRunbook.md",
            "project/docs/avalanche/avalanche-anchor-article.md",
            "project/docs/api/evidenceAnchoring.md",
        ],
    },
    {
        "number": 5,
        "name": "Pilot with enterprise and regulated AI use cases",
        "proof": "Signed Forg3t & HMA pilot agreement is linked and previewed as pilot evidence for regulated AI/compliance use cases.",
        "screenshot": "milestone-05-pilot-contract-drive-preview.png",
        "links": [("Signed pilot agreement", PILOT_CONTRACT_LINK)],
        "files": [
            "C:/Users/Alvinn/Downloads/Forg3t & HMA_signed.pdf",
            "project/docs/grant-review/submission-assets/screenshots/milestone-05-pilot-contract-drive-preview.png",
        ],
    },
    {
        "number": 6,
        "name": "Add role based admin workflows and reporting exports",
        "proof": "Reviewer job detail shows role-scoped workspace access, JSON export, evidence bundle, confirmed anchor, block number, and explorer link.",
        "screenshot": "milestone-06-rbac-reporting-exports.png",
        "links": [("Live app", LIVE_APP)],
        "files": [
            "project/src/lib/domainUtils.ts",
            "project/src/lib/domainUtils.test.ts",
            "project/src/pages/Settings.tsx",
            "project/supabase/functions/reports/index.ts",
            "project/supabase/functions/_shared/rbac.ts",
        ],
    },
    {
        "number": 7,
        "name": "Introduce repeatable verification pipelines for multiple unlearning jobs",
        "proof": "Phase 2 Review Pipeline 2026-06-04 produced 3 jobs, 3 evidence records, and 9 report exports; anchor coverage is demonstrated by the existing Avalanche mainnet anchor history in the same Phase 2 reviewer flow.",
        "screenshot": "milestone-07-repeatable-pipelines-anchor-history.png",
        "links": [("Live app", LIVE_APP)],
        "files": [
            "project/src/pages/Pipelines.tsx",
            "project/supabase/functions/pipelines/index.ts",
            "project/docs/grant-review/submission-assets/api/pipeline-review-run.json",
            "project/docs/grant-review/submission-assets/api/pipeline-review-list.json",
            "project/docs/grant-review/submission-assets/api/previous-mainnet-anchors.json",
        ],
    },
    {
        "number": 8,
        "name": "Expand integrations for teams using API based AI systems",
        "proof": "OpenAI-compatible black-box suppression flow and integrations settings support API-based AI teams with server-side secrets.",
        "screenshot": "milestone-08-openai-compatible-flow.png",
        "links": [("Live app", LIVE_APP)],
        "files": [
            "project/src/pages/Settings.tsx",
            "project/src/pages/Unlearning.tsx",
            "project/supabase/functions/integrations/index.ts",
            "project/docs/grant-review/evidence-artifacts/api/integration-create-generic-http.json",
            "project/docs/grant-review/evidence-artifacts/api/integration-test-generic-http.json",
        ],
    },
]


def ensure_dirs() -> None:
    for path in (VIDEO_DIR, API_DIR, RENDER_DIR):
        path.mkdir(parents=True, exist_ok=True)


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/segoeuib.ttf") if bold else Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf") if bold else Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def fit_image(canvas: Image.Image, image_path: Path, box: tuple[int, int, int, int]) -> None:
    image = Image.open(image_path).convert("RGB")
    x, y, width, height = box
    scale = min(width / image.width, height / image.height)
    size = (max(1, int(image.width * scale)), max(1, int(image.height * scale)))
    image = image.resize(size, Image.LANCZOS)
    canvas.paste(image, (x + (width - size[0]) // 2, y + (height - size[1]) // 2))


def wrap_text(draw: ImageDraw.ImageDraw, text: str, text_font, max_width: int) -> list[str]:
    lines: list[str] = []
    line = ""
    for word in text.split():
        candidate = f"{line} {word}".strip()
        if not line or draw.textbbox((0, 0), candidate, font=text_font)[2] <= max_width:
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def build_walkthrough_frame(title: str, subtitle: str, image_path: Path, step: int, total: int) -> Image.Image:
    width, height = 1280, 720
    frame = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(frame)
    draw.rectangle([0, 0, width, 86], fill="#F7F9FC")
    draw.line([0, 86, width, 86], fill="#DADCE0", width=2)
    draw.text((42, 24), "Forg3t Protocol - Avalanche Phase 2 Evidence Walkthrough", font=font(26, True), fill=f"#{INK}")
    draw.text((42, 55), title, font=font(17, True), fill=f"#{DARK_BLUE}")

    image_box = (42, 112, 1196, 498)
    draw.rounded_rectangle(
        [image_box[0] - 8, image_box[1] - 8, image_box[0] + image_box[2] + 8, image_box[1] + image_box[3] + 8],
        radius=12,
        outline="#CBD5E1",
        width=2,
        fill="#FFFFFF",
    )
    fit_image(frame, image_path, image_box)

    draw.rounded_rectangle([42, 634, 1238, 694], radius=10, fill="#F2F4F7", outline="#DADCE0")
    y = 650
    for line in wrap_text(draw, subtitle, font(22), 1130)[:2]:
        draw.text((64, y), line, font=font(22), fill=f"#{INK}")
        y += 27
    progress_width = int(1196 * ((step + 1) / total))
    draw.rectangle([42, 704, 42 + progress_width, 710], fill=f"#{BLUE}")
    return frame


def generate_pipeline_anchor_history_card() -> None:
    API_DIR.mkdir(parents=True, exist_ok=True)
    (API_DIR / "previous-mainnet-anchors.json").write_text(
        json.dumps({"anchors": PREVIOUS_MAINNET_ANCHORS}, indent=2),
        encoding="utf-8",
    )

    width, height = 1280, 720
    card = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(card)
    draw.rectangle([0, 0, width, 86], fill="#F7F9FC")
    draw.line([0, 86, width, 86], fill="#DADCE0", width=2)
    draw.text((42, 24), "Milestone 7 - Repeatable Verification Pipelines", font=font(28, True), fill=f"#{INK}")
    draw.text(
        (42, 58),
        "Pipeline automation plus existing Avalanche mainnet anchor history",
        font=font(17, True),
        fill=f"#{DARK_BLUE}",
    )

    pipeline_img = SS_DIR / "milestone-07-repeatable-pipelines.png"
    if pipeline_img.exists():
        draw.rounded_rectangle([42, 112, 610, 468], radius=12, fill="#FFFFFF", outline="#CBD5E1", width=2)
        image = Image.open(pipeline_img).convert("RGB")
        image = image.crop((0, 0, int(image.width * 0.62), image.height))
        image_path = SS_DIR / "_milestone-07-pipeline-crop.png"
        image.save(image_path)
        fit_image(card, image_path, (54, 124, 544, 332))
        image_path.unlink(missing_ok=True)
        draw.text((64, 480), "Production pipeline UI", font=font(18, True), fill=f"#{INK}")
        draw.text((64, 506), "Completed run creates repeatable jobs, evidence, and reports.", font=font(15), fill=f"#{MUTED}")

    draw.rounded_rectangle([650, 112, 1238, 468], radius=12, fill="#FFFFFF", outline="#CBD5E1", width=2)
    draw.text((674, 134), "Anchor history already available in the reviewer flow", font=font(20, True), fill=f"#{INK}")
    draw.text((674, 164), "These are prior Avalanche C-Chain mainnet anchors used by the same Phase 2 evidence flow.", font=font(14), fill=f"#{MUTED}")

    y = 204
    for item in PREVIOUS_MAINNET_ANCHORS[:5]:
        short_tx = f"{item['tx'][:12]}...{item['tx'][-8:]}"
        draw.rounded_rectangle([674, y, 1214, y + 46], radius=8, fill="#F8FAFC", outline="#E5E7EB")
        draw.text((692, y + 8), item["date"], font=font(15, True), fill=f"#{INK}")
        draw.text((806, y + 8), short_tx, font=font(15), fill=f"#{INK}")
        draw.text((806, y + 27), f"job {item['job'][:8]}... | block {item['block']}", font=font(12), fill=f"#{MUTED}")
        y += 54

    draw.rounded_rectangle([42, 534, 1238, 684], radius=12, fill="#F2F4F7", outline="#DADCE0")
    draw.text((64, 558), "Correct milestone interpretation", font=font(20, True), fill=f"#{INK}")
    lines = [
        "Latest pipeline run: 3 jobs, 3 evidence records, and 9 report exports.",
        "Anchor proof: existing confirmed mainnet anchors in the same Phase 2 reviewer evidence flow.",
        "The anchor history above is the mainnet proof trail used alongside the repeatable pipeline evidence.",
    ]
    y = 588
    for line in lines:
        draw.text((82, y), f"- {line}", font=font(17), fill=f"#{INK}")
        y += 28

    card.save(SS_DIR / "milestone-07-repeatable-pipelines-anchor-history.png")


def generate_contact_sheet() -> None:
    thumb_w, thumb_h = 560, 318
    label_h = 42
    sheet = Image.new("RGB", (thumb_w * 2 + 36, (thumb_h + label_h) * 4 + 54), "#FFFFFF")
    draw = ImageDraw.Draw(sheet)
    for index, item in enumerate(MILESTONES):
        image_path = SS_DIR / item["screenshot"]
        if not image_path.exists():
            continue
        image = Image.open(image_path).convert("RGB")
        image.thumbnail((thumb_w, thumb_h), Image.LANCZOS)
        x = 12 + (index % 2) * (thumb_w + 12)
        y = 12 + (index // 2) * (thumb_h + label_h + 10)
        draw.rounded_rectangle(
            [x, y, x + thumb_w, y + thumb_h + label_h],
            radius=8,
            outline="#D1D5DB",
            width=2,
            fill="#F9FAFB",
        )
        draw.text((x + 12, y + 9), f"M{item['number']} {item['name'][:32]}", fill=f"#{INK}", font=font(20))
        sheet.paste(image, (x + (thumb_w - image.width) // 2, y + label_h + (thumb_h - image.height) // 2))
    sheet.save(SS_DIR / "milestone-contact-sheet.png")


def generate_walkthrough() -> None:
    generate_pipeline_anchor_history_card()
    generate_contact_sheet()
    slides = [("Sign in", "Reviewer enters through the production buildgames.forg3t.io sign-in flow.", SS_DIR / "00-sign-in-start.png")]
    slides.extend(
        (
            f"M{milestone['number']} - {milestone['name']}",
            milestone["proof"],
            SS_DIR / milestone["screenshot"],
        )
        for milestone in MILESTONES
    )
    frames = [
        build_walkthrough_frame(title, subtitle, image_path, index, len(slides))
        for index, (title, subtitle, image_path) in enumerate(slides)
        if image_path.exists()
    ]
    if frames:
        frames[0].save(GIF_PATH, save_all=True, append_images=frames[1:], duration=1600, loop=0, optimize=True)
        frames[0].save(THUMB_PATH)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    props = cell._tc.get_or_add_tcPr()
    borders = props.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        props.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: float = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    style_paragraph(paragraph, after=0, line_spacing=1.08)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=24, color=INK, bold=True)


def add_subtitle(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=18)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=12, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    style_paragraph(paragraph, before=14 if level == 1 else 8, after=6 if level == 1 else 4, line_spacing=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color=BLUE if level == 1 else DARK_BLUE, bold=True)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=7, line_spacing=1.1)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        style_paragraph(paragraph, after=3, line_spacing=1.1)
        paragraph.paragraph_format.left_indent = Inches(0.35)
        run = paragraph.add_run(item)
        set_run_font(run, size=9.8, color=INK)


def add_link_line(doc: Document, label: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=3, line_spacing=1.1)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.5, color=INK, bold=True)
    add_hyperlink(paragraph, url, url)


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.35) -> None:
    if not image_path.exists():
        add_body(doc, f"Missing image: {image_path}")
        return
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, after=2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_p = doc.add_paragraph()
    style_paragraph(caption_p, after=9, line_spacing=1.0)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    set_run_font(run, size=8.2, color=MUTED, italic=True)


def add_simple_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell_shading(cell, LIGHT_FILL)
        cell_border(cell)
        cell_margins(cell)
        set_cell_text(cell, header, bold=True, color=INK, size=8.6)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell_shading(cells[index], "FFFFFF")
            cell_border(cells[index])
            cell_margins(cells[index])
            color = GREEN if index == 1 and value == "Completed" else INK
            set_cell_text(cells[index], value, bold=index == 0, color=color, size=8.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_metadata_rows(doc: Document, rows: list[tuple[str, str]]) -> None:
    for label, value in rows:
        paragraph = doc.add_paragraph()
        style_paragraph(paragraph, after=2, line_spacing=1.08)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10, color=INK)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Forg3t Protocol | Avalanche Grant Phase 2 Completion Packet | June 4, 2026")
    set_run_font(run, size=8, color=MUTED)
    return doc


def add_cover(doc: Document) -> None:
    add_title(doc, "Forg3t Protocol")
    add_subtitle(doc, "Avalanche Grant Phase 2 Completion Packet")
    add_body(
        doc,
        "This packet is prepared for Avalanche review and documents completion evidence across production Avalanche anchoring, auditor verification, job and evidence UX, documentation, pilot evidence, RBAC and exports, repeatable pipelines, and API-based AI integrations.",
    )
    add_metadata_rows(
        doc,
        [
            ("Prepared for", "Avalanche Grant Review"),
            ("Prepared date", "June 4, 2026"),
            ("Live app", LIVE_APP),
            ("GitHub", REPO_URL),
            ("Avalanche network", "C-Chain mainnet, chainId 43114"),
            ("Anchor contract", CONTRACT),
            ("Latest proof transaction", LATEST_TX_HASH),
            ("Latest proof block", LATEST_BLOCK),
        ],
    )
    add_link_line(doc, "GitHub repository", REPO_URL)
    add_link_line(doc, "Live production app", LIVE_APP)
    add_link_line(doc, "Medium architecture article", MEDIUM_ARTICLE)
    add_link_line(doc, "Signed pilot agreement", PILOT_CONTRACT_LINK)
    add_link_line(doc, "Latest public verify route", LATEST_VERIFY)
    add_link_line(doc, "Latest Snowtrace transaction", LATEST_TX)
    doc.add_page_break()


def add_summary(doc: Document) -> None:
    add_heading(doc, "Submission Readiness")
    add_body(
        doc,
        "All eight Avalanche Grant Phase 2 milestones are marked Completed in this packet. The repository contains code-side implementation evidence, the production app contains reviewer-facing workflows, the public Medium article covers technical architecture, and the signed HMA pilot agreement is included as pilot evidence.",
    )
    add_body(
        doc,
        "The Avalanche anchor proof is mainnet, not testnet: the C-Chain RPC returned chainId 43114, contract bytecode exists at the anchor contract address, the transaction receipt status is 0x1, and the same contract/transaction check on Fuji returned no contract bytecode and no receipt.",
    )
    add_simple_table(
        doc,
        ["Milestone", "Status", "Key proof"],
        [(f"{item['number']}. {item['name']}", "Completed", item["proof"]) for item in MILESTONES],
        [2.6, 1.0, 2.75],
    )
    add_image(doc, SS_DIR / "milestone-contact-sheet.png", "All milestone screenshots used in this packet.")
    doc.add_page_break()


def add_mainnet_verification(doc: Document) -> None:
    add_heading(doc, "Mainnet Contract Verification")
    verification_path = API_DIR / "avalanche-mainnet-verification.json"
    if verification_path.exists():
        verification = json.loads(verification_path.read_text(encoding="utf-8-sig"))
        mainnet = verification["mainnet"]
        fuji = verification["fujiControl"]
        add_simple_table(
            doc,
            ["Check", "Result"],
            [
                ("Mainnet RPC", mainnet["rpc"]),
                ("Mainnet chainId", f"{mainnet['chainIdDecimal']} ({mainnet['chainIdHex']})"),
                ("Contract bytecode", f"{mainnet['contractCodeBytesApprox']} bytes approx."),
                ("Receipt status", mainnet["receiptStatus"]),
                ("Block number", str(mainnet["receiptBlockDecimal"])),
                ("Transaction to", mainnet["transactionTo"]),
                ("Fuji control", f"chainId {fuji['chainIdDecimal']}; code bytes {fuji['contractCodeBytesApprox']}; tx receipt null {fuji['transactionReceiptIsNull']}"),
            ],
            [2.0, 4.35],
        )
    add_image(doc, SS_DIR / "milestone-01-avalanche-mainnet-anchor.png", "Milestone 1 proof: Snowtrace mainnet transaction details.")
    add_link_line(doc, "Snowtrace transaction", LATEST_TX)
    doc.add_page_break()


def add_milestone_sections(doc: Document) -> None:
    for item in MILESTONES:
        add_heading(doc, f"Milestone {item['number']}: {item['name']}")
        add_body(doc, "Status: Completed.")
        add_body(doc, item["proof"])
        add_image(doc, SS_DIR / item["screenshot"], f"Milestone {item['number']} evidence screenshot.")
        if item["links"]:
            for label, url in item["links"]:
                add_link_line(doc, label, url)
        add_body(doc, "Repository and evidence artifacts:")
        add_bullets(doc, item["files"])
        if item["number"] in {2, 4, 6}:
            doc.add_page_break()


def add_final_package(doc: Document) -> None:
    add_heading(doc, "Final Readiness Statement")
    add_body(
        doc,
        "Forg3t Protocol is ready to submit the Avalanche Grant Phase 2 disbursement request. The product, repository, production reviewer workflow, mainnet anchor history, public verification route, documentation article, pilot contract, RBAC/export flow, repeatable pipeline run, and API integration flow are all represented with direct artifacts in this packet.",
    )
    add_simple_table(
        doc,
        ["Artifact", "Location"],
        [
            ("Word packet", "project/docs/grant-review/forg3t-avalanche-phase2-submission-packet.docx"),
            ("Walkthrough GIF", "project/docs/grant-review/submission-assets/video/forg3t-avalanche-phase2-proof-walkthrough.gif"),
            ("Milestone screenshots", "project/docs/grant-review/submission-assets/screenshots/"),
            ("Mainnet verification JSON", "project/docs/grant-review/submission-assets/api/avalanche-mainnet-verification.json"),
            ("Pipeline run JSON", "project/docs/grant-review/submission-assets/api/pipeline-review-run.json"),
            ("Medium article", MEDIUM_ARTICLE),
            ("Signed pilot agreement", PILOT_CONTRACT_LINK),
        ],
        [2.0, 4.35],
    )
    add_body(
        doc,
        "Reviewer account credentials should be shared with Avalanche through a secure channel only. Passwords, admin database credentials, and wallet signer secrets are intentionally not embedded in this document or committed to the repository.",
    )


def build_docx() -> None:
    doc = setup_document()
    add_cover(doc)
    add_summary(doc)
    add_mainnet_verification(doc)
    add_milestone_sections(doc)
    add_final_package(doc)
    doc.core_properties.title = "Forg3t Protocol Avalanche Grant Phase 2 Completion Packet"
    doc.core_properties.subject = "Avalanche Phase 2 milestone evidence package"
    doc.core_properties.author = "Forg3t Protocol"
    doc.save(DOCX_PATH)


def audit_docx() -> dict[str, int]:
    with ZipFile(DOCX_PATH) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
    doc = Document(DOCX_PATH)
    return {
        "docx_bytes": DOCX_PATH.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "media_files": len(media_files),
        "gif_bytes": GIF_PATH.stat().st_size if GIF_PATH.exists() else 0,
    }


def main() -> None:
    ensure_dirs()
    generate_walkthrough()
    build_docx()
    result = {
        "docx": str(DOCX_PATH),
        "gif": str(GIF_PATH),
        "thumbnail": str(THUMB_PATH),
        **audit_docx(),
    }
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
