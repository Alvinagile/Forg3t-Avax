from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "grant-review"
OUT_DOCX = OUT_DIR / "forg3t-avalanche-phase2-grant-manager-packet.docx"

APP_URL = "https://buildgames.forg3t.io"
PUBLIC_VERIFY_URL = "https://buildgames.forg3t.io/verify/212a028eabcac5e2c3458cc9f219259d"
SUPABASE_VERIFY_URL = (
    "https://xewxfsdrtqpthkpbhbzp.supabase.co/functions/v1/"
    "verify-evidence?token=212a028eabcac5e2c3458cc9f219259d"
)
SNOWTRACE_TX_URL = (
    "https://snowtrace.io/tx/"
    "0x7af8b0376079571f2a4ff46ff76e6cdfb27f710ea4b2434c41bfdf25a167e7be"
)
SNOWTRACE_CONTRACT_URL = "https://snowtrace.io/address/0x20E772a60CEE7D8E6706E698B129FD917c3936bf"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B57D0")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (2.0, 4.2)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color="1F2937")
        set_cell_text(cells[1], value)
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
    return table


def add_status_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "1F2937")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    return table


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_link_line(doc: Document, label: str, url: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"{label}: ").bold = True
    add_hyperlink(p, url, url)
    return p


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 16, "E84142"),
        ("Heading 2", 13, "1F2937"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Forg3t Protocol Avalanche Grant Phase 2")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("111827")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Grant Manager Review Packet | Code-Side Evidence and Live Demo Access").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared: 2026-05-24 | Network: Avalanche C-Chain mainnet | Status: code-side review ready")

    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "Forg3t Protocol now has a reviewable Phase 2 product implementation for black-box AI suppression/unlearning evidence workflows. "
        "The repository-controlled milestones have been implemented or strengthened across the frontend, Supabase backend, database/RBAC, "
        "Avalanche anchoring, reports, pipelines, API integrations, tests, documentation, and dependency audit posture."
    )
    doc.add_paragraph(
        "This packet intentionally does not claim enterprise pilot approval, customer attestation, real customer usage, founder sales material, "
        "or recorded demo completion. Those items require external human evidence and should be supplied separately."
    )

    add_heading(doc, "Live Review Access", 1)
    add_kv_table(
        doc,
        [
            ("Production app", APP_URL),
            ("Reviewer email", "grant-reviewer@forg3t.io"),
            ("Reviewer password", "Share separately via a secure channel; do not commit or embed."),
            ("Reviewer role", "admin on the Phase 2 review project"),
            ("Public verify route", PUBLIC_VERIFY_URL),
            ("Public API verify", SUPABASE_VERIFY_URL),
            ("Snowtrace transaction", SNOWTRACE_TX_URL),
        ],
    )
    doc.add_paragraph(
        "Security note: this reviewer account is intended only for grant review. Rotate or delete it after the review window."
    )

    add_heading(doc, "Primary Evidence Record", 1)
    add_kv_table(
        doc,
        [
            ("Job ID", "9a2a77cf-4b09-4f59-be04-c18b08b137bd"),
            ("Evidence ID", "608a427f-25c1-4f43-b3e7-d9a86ff33801"),
            ("Evidence hash", "0x1c01887f16bc967d3e53042db25c8cfe2ad9611222a1b7a2f8326190061b4b0b"),
            ("Job hash", "0x9b7077f20c7a481b176684181bc7c28cab90a96d93ce73463dccfe12ccef649e"),
            ("Anchor status", "confirmed"),
            ("Network", "mainnet"),
            ("Chain ID", "43114"),
            ("Block number", "86276105"),
            ("Contract", "0x20E772a60CEE7D8E6706E698B129FD917c3936bf"),
            ("Transaction", "0x7af8b0376079571f2a4ff46ff76e6cdfb27f710ea4b2434c41bfdf25a167e7be"),
        ],
        widths=(1.55, 4.75),
    )

    add_heading(doc, "Milestone Completion Matrix", 1)
    add_status_table(
        doc,
        ["Milestone", "Status", "Evidence"],
        [
            [
                "Evidence anchoring",
                "Completed",
                "Job to evidence hash to Avalanche mainnet anchor to verification result. Tx hash, network, block, contract, and Snowtrace link are visible.",
            ],
            [
                "Drag-and-drop verification",
                "Completed",
                "Dashboard verify flow supports JSON/PDF-style evidence checks and states: valid, mismatch, pending, confirmed, failed.",
            ],
            [
                "Job history and evidence UX",
                "Completed",
                "Job list/detail, evidence detail, public verify links, anchor visibility, and report access were strengthened for reviewer navigation.",
            ],
            [
                "Reporting exports and RBAC",
                "Completed",
                "JSON, CSV, and PDF exports are available from job/evidence records. RBAC behavior covers owner, admin, compliance, auditor, developer, viewer.",
            ],
            [
                "Repeatable pipelines",
                "Completed",
                "Pipeline runs can create multiple jobs and move them through evidence generation, optional anchoring, verification, and report export where configured.",
            ],
            [
                "API-based AI integrations",
                "Completed",
                "OpenAI-compatible and generic HTTP integration flows are documented with lifecycle examples from setup to job/evidence retrieval.",
            ],
            [
                "Documentation",
                "Completed",
                "README and Phase 2 readiness docs include commands, routes, expected outputs, and non-code exclusions.",
            ],
            [
                "Dependency audit",
                "Completed",
                "Frontend package and contracts package both return npm audit: 0 vulnerabilities.",
            ],
        ],
        widths=[1.55, 1.0, 3.85],
    )

    add_heading(doc, "Reviewer Walkthrough", 1)
    for step in [
        "Open the production app and sign in with the reviewer account above.",
        "Go to Jobs and open the completed Phase 2 smoke job.",
        "Confirm evidence hash, job hash, anchor status, network, block number, transaction hash, and explorer link.",
        "Open the Evidence detail page and review manifest/report payload/export access.",
        "Open the public verify route and confirm Anchor Confirmed on mainnet.",
        "Open the Snowtrace transaction link and confirm the C-Chain transaction exists.",
        "Use Verify to test evidence JSON/PDF drag-and-drop verification states.",
        "Open Pipelines to inspect repeatable pipeline run support.",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "Review Links", 1)
    add_link_line(doc, "Production app", APP_URL)
    add_link_line(doc, "Public evidence verification page", PUBLIC_VERIFY_URL)
    add_link_line(doc, "Public verify Edge Function", SUPABASE_VERIFY_URL)
    add_link_line(doc, "Avalanche mainnet transaction", SNOWTRACE_TX_URL)
    add_link_line(doc, "Avalanche contract address", SNOWTRACE_CONTRACT_URL)

    add_heading(doc, "Commands and Verification Outputs", 1)
    add_status_table(
        doc,
        ["Command", "Result"],
        [
            ("npm audit", "0 vulnerabilities in project package"),
            ("npm test", "4 test files passed, 15 tests passed"),
            ("npm run lint", "Passed"),
            ("npm run build", "Passed with Vite 8 production build"),
            ("cd contracts; npm audit", "0 vulnerabilities in contracts package"),
            ("cd contracts; npm run compile", "Passed with Node 22.13+ / Node 24 runtime"),
            ("contracts local deploy smoke", "Passed on local Hardhat simulated network"),
        ],
        widths=[2.7, 3.7],
    )

    add_heading(doc, "Repository Evidence Files", 1)
    for item in [
        "README.md and project/README.md",
        "project/docs/phase2-readiness.md",
        "project/docs/api/evidenceAnchoring.md",
        "project/scripts/phase2-smoke.mjs",
        "project/supabase/functions/anchors/index.ts",
        "project/supabase/functions/verify-evidence/index.ts",
        "project/supabase/functions/jobs/index.ts",
        "project/supabase/functions/pipelines/index.ts",
        "project/supabase/functions/reports/index.ts",
        "project/supabase/functions/integrations/index.ts",
        "project/contracts/contracts/ForgEvidenceAnchor.sol",
        "project/contracts/hardhat.config.ts",
        "project/shared/workflows.test.ts and project/src/lib/domainUtils.test.ts",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Explicit Non-Code Exclusions", 1)
    doc.add_paragraph(
        "The following items are not claimed as completed in this repository packet because they require external proof or manual review artifacts:"
    )
    for item in [
        "Enterprise pilot approval",
        "Customer attestation",
        "Real customer usage evidence",
        "Founder/sales material",
        "Recorded demo/video",
        "Optional Snowtrace contract source verification, if the reviewer requests it separately",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Known Review Risks", 1)
    for item in [
        "Black-box suppression verification is supported; internal model-weight deletion is not claimed unless an integration supplies verifiable internal evidence.",
        "Contracts now use Hardhat 3 to clear audit findings, which requires Node 22.13.0 or newer.",
        "Hardhat automatic Snowtrace source verification plugin was removed because it reintroduced audited legacy ethers dependencies; manual Snowtrace verification remains possible.",
        "The reviewer account is live and should be rotated after review.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Final Code-Side Readiness Statement", 1)
    doc.add_paragraph(
        "All repository-controlled Phase 2 product and code milestones are implemented or strengthened. "
        "Remaining items are external proof, human approval, demo capture, or optional manual contract-source verification."
    )

    doc.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    print(build())
